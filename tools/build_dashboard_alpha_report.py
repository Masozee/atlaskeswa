from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "reports" / "dashboard_ommha"
SHOTS = BASE / "alpha_screenshots"
OUT = BASE / "Frontend_Dashboard_OMMHA_Alpha_Development_and_Functional_Test_Report.docx"
LOGO = ROOT / "frontend" / "public" / "OMMHA.png"

BLUE = "07579E"
PALE_BLUE = "EAF3F8"
GRAY = "666666"
LIGHT_GRAY = "F4F7FA"
GREEN = "2E7D32"
AMBER = "9A6700"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    element = tc_pr.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        tc_pr.append(element)
    element.set(qn("w:fill"), fill)


def cell_text(cell, value, bold=False, color=None, size=8.2):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, value=85):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Frontend Dashboard OMMHA — Alpha Version   |   ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Update this field in Microsoft Word to display the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separator, placeholder, end])


def body(doc, text, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.italic = italic
    return paragraph


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, rows, widths=None):
    result = doc.add_table(rows=1, cols=len(rows[0]))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(rows[0]):
        cell_text(result.rows[0].cells[index], value, bold=True, color="FFFFFF", size=8.4)
        shade(result.rows[0].cells[index], BLUE)
    for row_index, row in enumerate(rows[1:], start=1):
        cells = result.add_row().cells
        for column_index, value in enumerate(row):
            cell_text(cells[column_index], value)
            set_cell_margins(cells[column_index])
            if row_index % 2 == 0:
                shade(cells[column_index], LIGHT_GRAY)
    if widths:
        for row in result.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return result


def figure(doc, filename, caption, width=6.45):
    path = SHOTS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.style = doc.styles["Caption"]
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def section(doc, title, subtitle=None):
    doc.add_page_break()
    heading = doc.add_heading(title, level=1)
    if subtitle:
        paragraph = doc.add_paragraph(subtitle)
        paragraph.style = doc.styles["Subtitle"]


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (
        ("Title", 24, BLUE),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, "222222"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)

    doc.styles["Subtitle"].font.name = "Arial"
    doc.styles["Subtitle"].font.size = Pt(11)
    doc.styles["Subtitle"].font.color.rgb = RGBColor.from_string(GRAY)
    doc.styles["Caption"].font.name = "Arial"
    doc.styles["Caption"].font.size = Pt(8.5)
    doc.styles["Caption"].font.italic = True
    doc.styles["Caption"].font.color.rgb = RGBColor.from_string(GRAY)

    header = sec.header.paragraphs[0]
    header.text = "OMMHA  |  One Map for Mental Health Atlas"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_page_number(sec.footer.paragraphs[0])

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)

    core = doc.core_properties
    core.title = "Frontend Dashboard OMMHA Alpha Development and Functional Test Report"
    core.subject = "Historical frontend snapshot and non-destructive functional verification"
    core.author = "Pusat Rehabilitasi YAKKUM"
    core.keywords = "OMMHA, Alpha, frontend, dashboard, functional testing"
    return doc


def build():
    doc = setup_document()

    for _ in range(2):
        doc.add_paragraph()
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run().add_picture(str(LOGO), width=Inches(1.55))
    for text, size in (
        ("DEVELOPMENT AND FUNCTIONAL TEST REPORT", 19),
        ("OMMHA FRONTEND DASHBOARD", 25),
    ):
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.name = "Arial"
        paragraph.runs[0].font.size = Pt(size)
        paragraph.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = doc.add_paragraph(
        "Dashboard Feature Documentation for Managing and Analysing Data\n"
        "for the Kebumen Mental Health Service Atlas"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = "Arial"
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph()
    version = doc.add_paragraph("Frontend Dashboard OMMHA — Alpha Version")
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].bold = True
    version.runs[0].font.size = Pt(12)
    date = doc.add_paragraph("Historical baseline: 29 June 2026")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.size = Pt(11)
    revision = doc.add_paragraph("Repository revision: fbc83cc")
    revision.alignment = WD_ALIGN_PARAGRAPH.CENTER
    revision.runs[0].font.size = Pt(9)
    revision.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    for _ in range(4):
        doc.add_paragraph()
    owner = doc.add_paragraph("Pusat Rehabilitasi YAKKUM\nOne Map for Mental Health Atlas (OMMHA)")
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owner.runs[0].font.name = "Arial"
    owner.runs[0].bold = True
    owner.runs[0].font.size = Pt(10.5)

    section(doc, "Table of Contents")
    add_toc(doc.add_paragraph())
    body(
        doc,
        "This report is an English, historical Alpha-version counterpart to the later Beta report. "
        "It is based on repository evidence and screenshots captured from an isolated checkout.",
    )

    section(doc, "A. Executive Summary")
    body(
        doc,
        "This report documents the OMMHA web frontend as it existed at revision fbc83cc on "
        "29 June 2026. This revision is the direct predecessor of the 8 July 2026 dashboard "
        "redesign that introduced a blue background across the desktop sidebar. The Alpha "
        "interface shown here retains the light sidebar identified as the historical boundary.",
    )
    body(
        doc,
        "The historical frontend was launched from an isolated archive on port 3100 and connected "
        "to the local YAKKUM backend on port 8001. The current source tree, the existing Beta report, "
        "and application records were not edited. Testing was limited to authentication, route "
        "loading, navigation, and read-only interface interactions.",
    )
    figure(doc, "00-login.png", "Figure 1. Alpha-version OMMHA login page")

    section(doc, "B. Purpose and Scope")
    doc.add_heading("1. Purpose", level=2)
    numbered(
        doc,
        [
            "Identify the dashboard functions present in the Alpha baseline.",
            "Preserve visual evidence of the light-sidebar interface.",
            "Verify that representative frontend routes render against a local backend.",
            "Record the boundary between verified behaviour and functions not exercised.",
            "Provide an auditable historical counterpart to the Beta-version report.",
        ],
    )
    doc.add_heading("2. Scope", level=2)
    body(
        doc,
        "The scope is the web frontend reachable through the authenticated dashboard navigation. "
        "The backend, database, mobile applications, deployment infrastructure, and production "
        "environment are integration context only and are not independently certified by this report.",
    )
    table(
        doc,
        [
            ["Functional group", "Documented Alpha coverage"],
            ["Dashboard", "Summary, indicators, map, submissions, and verification entry points"],
            ["Survey Management", "Records, new entry, status views, bulk upload, models, templates, and audit"],
            ["Service Management", "Service lists, creation entry point, and classification references"],
            ["Users and Roles", "Accounts, roles, login history, enumerator functions"],
            ["Data, Maps, and Reports", "Tables, matrices, coverage, gaps, geospatial views, and exports"],
            ["System and Help", "Settings, logs, user guidance, DESDE-LTC reference, FAQ, and support"],
        ],
        [4.3, 11.7],
    )

    section(doc, "C. Historical Baseline and Safety Controls")
    table(
        doc,
        [
            ["Item", "Value"],
            ["Alpha source revision", "fbc83cc920625c84ed4afaf9cae80bf4ca4a5957"],
            ["Revision date", "29 June 2026, 08:11:59 UTC+07:00"],
            ["Reason for boundary", "Direct predecessor of the full blue-sidebar change in a395566"],
            ["Screenshot date", "24 July 2026"],
            ["Historical frontend", "Isolated archive, http://127.0.0.1:3100"],
            ["Local YAKKUM backend", "http://127.0.0.1:8001"],
            ["Viewport", "1440 × 900 pixels"],
            ["Mutation boundary", "No create, edit, delete, approve, import, or export action executed"],
        ],
        [4.5, 11.5],
    )
    body(
        doc,
        "Because screenshots were captured on 24 July, live dates and aggregate values visible in "
        "the interface reflect the local database at capture time. The component structure, routes, "
        "navigation, and light-sidebar styling come from the 29 June Alpha revision.",
        italic=True,
    )

    section(doc, "D. Alpha Interface Overview")
    body(
        doc,
        "The Alpha desktop layout uses a narrow icon rail, a light secondary navigation panel, and "
        "a wide content area. Selecting a primary icon changes the submenu group. Breadcrumbs and "
        "page headings identify the current context, while cards, tables, forms, graphs, and maps "
        "occupy the main workspace.",
    )
    figure(doc, "01-dashboard.png", "Figure 2. Alpha dashboard summary with the light sidebar")
    bullets(
        doc,
        [
            "The left navigation background is light rather than the later Beta blue.",
            "Summary cards show services, pending surveys, active users, and district coverage.",
            "The activity overview and calendar provide time-oriented monitoring.",
            "Role-aware navigation exposes operational and administrative groups.",
        ],
    )

    section(doc, "E. Dashboard and Indicators")
    body(
        doc,
        "The summary is the post-login landing page. The Indicators page expands the overview into "
        "analytical cards and charts intended to support programme monitoring and comparison.",
    )
    figure(doc, "02-indicators.png", "Figure 3. Alpha main indicators page")
    table(
        doc,
        [
            ["Visible capability", "Alpha purpose"],
            ["Metric cards", "Provide an immediate view of aggregate values"],
            ["Activity chart", "Show recent activity over time"],
            ["Calendar", "Relate activity and submissions to dates"],
            ["Analytical visualisations", "Summarise distribution and coverage"],
            ["Sidebar navigation", "Move between dashboard, map, and submission contexts"],
        ],
        [5.1, 10.9],
    )

    section(doc, "F. Survey Management", "Primary documented workflow")
    body(
        doc,
        "Survey Management is the most detailed Alpha module. It links captured records to status "
        "views, questionnaire models, templates, and audit information.",
    )
    figure(doc, "03-surveys.png", "Figure 4. Alpha list of all survey records")
    bullets(
        doc,
        [
            "The table presents facility, geographic, questionnaire, and status information.",
            "Category, date, district, village, and enumerator filters are visible.",
            "The category filter was opened successfully during read-only verification.",
            "The sorting control was visible, but its expected accessible name could not be located; this check is Limited.",
        ],
    )
    doc.add_heading("1. New survey entry", level=2)
    body(
        doc,
        "The new-entry route rendered the historical questionnaire workflow. No field was submitted "
        "and no draft record was created during verification.",
    )
    figure(doc, "04-new-survey.png", "Figure 5. Alpha new-survey entry route")
    doc.add_heading("2. Pending submissions", level=2)
    body(
        doc,
        "The Pending Submissions view isolates records awaiting review. The page loaded and rendered; "
        "approval, rejection, and record-editing actions were deliberately excluded.",
    )
    figure(doc, "05-pending.png", "Figure 6. Alpha pending-submissions view")
    doc.add_heading("3. Bulk upload", level=2)
    body(
        doc,
        "The bulk-upload page provides a file-based intake path. The interface rendered, but no file "
        "was selected or uploaded.",
    )
    figure(doc, "06-bulk-upload.png", "Figure 7. Alpha bulk-upload interface")

    section(doc, "G. Questionnaire Models, Templates, and Audit")
    body(
        doc,
        "Questionnaire models define survey structure. During testing, the model page loaded and its "
        "search input accepted text without submitting a change.",
    )
    figure(doc, "07-question-models.png", "Figure 8. Alpha questionnaire-model list")
    body(
        doc,
        "The template route redirected to the available template detail and rendered the question map "
        "for OMMHA_V1 version 1.0. This verifies route resolution and visualisation rendering only.",
    )
    figure(doc, "08-templates.png", "Figure 9. Alpha survey-template question map")
    body(
        doc,
        "The survey audit page rendered historical activity information. No audit entry was modified.",
    )
    figure(doc, "09-audit.png", "Figure 10. Alpha survey audit log")

    section(doc, "H. Supporting Modules")
    doc.add_heading("1. Service Management", level=2)
    body(
        doc,
        "Service Management provides access to service records and reference classifications, including "
        "BSIC, MTC, service types, geographic units, and target populations.",
    )
    figure(doc, "10-services.png", "Figure 11. Alpha Service Management page")
    doc.add_heading("2. Map and Geospatial Functions", level=2)
    body(
        doc,
        "The map route rendered and the map zoom control responded. Heatmap, MTC-layer, regional "
        "comparison, and geospatial-upload routes were identified from the Alpha navigation but were "
        "not used to mutate or upload data.",
    )
    figure(doc, "11-map.png", "Figure 12. Alpha service-location map")
    doc.add_heading("3. Users and Roles", level=2)
    body(
        doc,
        "The users page rendered account and role information. Account creation, role changes, password "
        "changes, and deactivation were excluded from testing.",
    )
    figure(doc, "12-users.png", "Figure 13. Alpha user-management page")
    doc.add_heading("4. Reports and Export", level=2)
    body(
        doc,
        "The report-export page rendered its format and scope controls. No export was initiated because "
        "that action can create files and activity-log records.",
    )
    figure(doc, "13-export.png", "Figure 14. Alpha report and export interface")
    doc.add_heading("5. Help and Documentation", level=2)
    body(
        doc,
        "The embedded user guide rendered successfully. The Alpha navigation also exposes DESDE-LTC "
        "reference material, enumerator guidance, frequently asked questions, and technical support.",
    )
    figure(doc, "14-help.png", "Figure 15. Alpha embedded user guide")

    section(doc, "I. Functional Verification Results")
    body(
        doc,
        "Fourteen authenticated routes were loaded after a successful login. Each returned HTTP 200, "
        "remained on the intended application route (except the expected template-detail redirect), "
        "and produced no recorded frontend page error or failed local API response during capture.",
    )
    table(
        doc,
        [
            ["Check", "Result", "Evidence/constraint"],
            ["Login and redirect", "PASS", "Existing account authenticated and reached /dashboard"],
            ["Dashboard summary", "PASS", "Rendered with live cards, chart, calendar, and light sidebar"],
            ["Sidebar navigation", "PASS", "Navigation to Indicators completed"],
            ["Indicators", "PASS", "Route rendered without recorded page/API errors"],
            ["Survey list", "PASS", "Table and filter controls rendered with data"],
            ["Survey category filter", "PASS", "Filter menu opened"],
            ["Survey sorting interaction", "LIMITED", "Expected accessible control name not located"],
            ["New survey route", "PASS", "Form route rendered; no submission attempted"],
            ["Pending surveys", "PASS", "Route rendered; no approval/rejection attempted"],
            ["Bulk upload", "PASS", "Interface rendered; no file uploaded"],
            ["Question-model search", "PASS", "Search input accepted text"],
            ["Template visualisation", "PASS", "Redirected to template 1 and rendered question map"],
            ["Survey audit", "PASS", "Route rendered; no record changed"],
            ["Services", "PASS", "List route rendered with local data"],
            ["Map zoom", "PASS", "Map rendered and zoom control responded"],
            ["Users", "PASS", "User list rendered; no account action attempted"],
            ["Export interface", "PASS", "Route rendered; no download initiated"],
            ["Help content", "PASS", "User-guide content rendered"],
        ],
        [5.0, 2.2, 8.8],
    )

    section(doc, "J. Limitations and Interpretation")
    bullets(
        doc,
        [
            "PASS means the historical frontend route or named read-only interaction worked in the isolated local test; it is not a production certification.",
            "The backend and database were current local copies, not a full database snapshot from 29 June 2026.",
            "Authentication naturally updates security metadata such as login time and token records; no domain record was intentionally created or changed.",
            "Create, edit, delete, upload, import, export, approve, reject, password, permission, and settings changes were not exercised.",
            "Values and dates visible in screenshots reflect the local dataset at capture time.",
            "Responsive/mobile browser layouts, performance, security, accessibility conformance, and cross-browser compatibility were outside this report.",
        ],
    )

    section(doc, "Appendix 1. Alpha Frontend Feature Matrix")
    table(
        doc,
        [
            ["No.", "Group", "Page or feature", "Historical evidence"],
            ["1", "Dashboard", "Summary", "Route and screenshot verified"],
            ["2", "Dashboard", "Main Indicators", "Route, navigation, and screenshot verified"],
            ["3", "Dashboard", "Service Distribution Map", "Present in Alpha navigation"],
            ["4", "Dashboard", "Recent Submissions", "Present in Alpha navigation"],
            ["5", "Dashboard", "Verification Queue", "Present in Alpha navigation"],
            ["6", "Survey", "All Survey Records", "Route, data table, filter, and screenshot verified"],
            ["7", "Survey", "New Survey Entry", "Route and screenshot verified; no submission"],
            ["8", "Survey", "Pending Submissions", "Route and screenshot verified"],
            ["9", "Survey", "Approved Surveys", "Present in Alpha navigation"],
            ["10", "Survey", "Rejected Surveys", "Present in Alpha navigation"],
            ["11", "Survey", "Bulk Upload", "Route and screenshot verified; no upload"],
            ["12", "Survey", "Questionnaire Models", "Route, search input, and screenshot verified"],
            ["13", "Survey", "Survey Templates", "Detail redirect, visualisation, and screenshot verified"],
            ["14", "Survey", "Survey Audit Log", "Route and screenshot verified"],
            ["15", "Services", "All Services", "Route and screenshot verified"],
            ["16", "Services", "Reference Data", "BSIC, MTC, types, geography, and target population in navigation"],
            ["17", "Users", "Users and Roles", "User-list route and screenshot verified"],
            ["18", "Data", "Tables and Matrices", "Present in Alpha navigation"],
            ["19", "Map", "Service Locations", "Route, zoom interaction, and screenshot verified"],
            ["20", "Reports", "Report and Export", "Route and screenshot verified; no export"],
            ["21", "System", "Settings and Logs", "Present in Alpha navigation"],
            ["22", "Help", "User Guide", "Route, content, and screenshot verified"],
        ],
        [1.0, 2.5, 5.3, 7.2],
    )

    section(doc, "K. Conclusion")
    body(
        doc,
        "The OMMHA Alpha frontend at revision fbc83cc provided an integrated dashboard for survey "
        "management, service information, analytics, maps, user administration, reporting, system "
        "monitoring, and embedded help. Its distinguishing visual characteristic is the light desktop "
        "sidebar that preceded the full blue-sidebar redesign.",
    )
    body(
        doc,
        "The isolated historical build authenticated and rendered all representative routes selected "
        "for this report. Read-only navigation and interaction checks were successful except for the "
        "survey sorting control, which remains Limited. No application-code rollback was performed, "
        "and the current Beta report was preserved.",
    )
    ending = doc.add_paragraph("— End of report —")
    ending.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ending.runs[0].bold = True
    ending.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    BASE.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
