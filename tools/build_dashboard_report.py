from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "reports" / "dashboard_ommha"
SHOTS = BASE / "screenshots"
OUT = BASE / "Laporan_Pengembangan_dan_Uji_Fungsional_Frontend_Dashboard_OMMHA_Beta.docx"
LOGO = ROOT / "frontend" / "public" / "OMMHA.png"

TEAL = "07579E"
LIGHT_TEAL = "EAF3F8"
GOLD = "F2B84B"
GRAY = "666666"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Frontend Dashboard OMMHA Versi Beta   |   ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Daftar isi akan diperbarui otomatis saat dokumen dibuka."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def page_break(doc):
    doc.add_page_break()


def add_title(doc, title, subtitle=None):
    heading = doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.style = doc.styles["Subtitle"]
    return heading


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_start and text.startswith(bold_start):
        p.add_run(bold_start).bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_figure(doc, filename, caption, width=6.45):
    path = SHOTS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph(caption)
    cp.style = doc.styles["Caption"]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_feature_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], value, bold=True, color="FFFFFF", size=8.5)
        shade(table.rows[0].cells[i], TEAL)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=8.2)
            if len(table.rows) % 2 == 1:
                shade(cells[i], "F4F7FA")
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def new_page_section(doc, title, subtitle=None):
    heading = add_title(doc, title, subtitle)
    heading.paragraph_format.page_break_before = True


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (("Title", 24, TEAL), ("Heading 1", 17, TEAL), ("Heading 2", 13, TEAL), ("Heading 3", 11, "222222")):
        s = doc.styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after = Pt(6)

    doc.styles["Subtitle"].font.name = "Arial"
    doc.styles["Subtitle"].font.size = Pt(11)
    doc.styles["Subtitle"].font.color.rgb = RGBColor.from_string(GRAY)
    doc.styles["Caption"].font.name = "Arial"
    doc.styles["Caption"].font.size = Pt(8.5)
    doc.styles["Caption"].font.italic = True
    doc.styles["Caption"].font.color.rgb = RGBColor.from_string(GRAY)

    for sec in doc.sections:
        h = sec.header.paragraphs[0]
        h.text = "OMMHA  |  One Map for Mental Health Atlas"
        h.runs[0].font.name = "Arial"
        h.runs[0].font.size = Pt(8)
        h.runs[0].font.bold = True
        h.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
        add_page_number(sec.footer.paragraphs[0])

    settings = doc.settings._element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    return doc


def build():
    doc = setup_document()

    # 1 — Cover
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(LOGO), width=Inches(1.55))
    p = doc.add_paragraph("LAPORAN PENGEMBANGAN DAN UJI FUNGSIONAL")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(20)
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph("FRONTEND DASHBOARD OMMHA")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(25)
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph("Dokumentasi Fitur Dashboard untuk Mendukung Pengelolaan dan Analisis Data\nAtlas Layanan Kesehatan Jiwa Kabupaten Kebumen")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph()
    p = doc.add_paragraph("Frontend Dashboard OMMHA Versi Beta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    p = doc.add_paragraph("20 Juli 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(11)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph("Pusat Rehabilitasi YAKKUM\nOne Map for Mental Health Atlas (OMMHA)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "Arial"
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10.5)

    # 2 — TOC
    new_page_section(doc, "Daftar Isi")
    add_toc(doc.add_paragraph())
    doc.add_paragraph()
    add_body(doc, "Dokumen ini disusun dalam tujuh bagian utama dan lampiran matriks fitur. Daftar isi dapat diperbarui melalui perintah Update Field di Microsoft Word.")

    # 3 — Background
    new_page_section(doc, "A. Latar Belakang")
    add_body(doc, "One Map for Mental Health Atlas (OMMHA) dikembangkan untuk mendukung pemetaan ketersediaan, jenis, kapasitas, dan distribusi layanan kesehatan jiwa di Kabupaten Kebumen. Informasi tersebut diperlukan agar pemangku kepentingan dapat melihat layanan bukan hanya berdasarkan nama institusi, tetapi juga berdasarkan jenis layanan yang benar-benar diberikan kepada pengguna layanan.")
    add_body(doc, "Kerangka Description and Evaluation of Services and Directories in Europe for Long-Term Care (DESDE-LTC) digunakan sebagai dasar klasifikasi terstandar. Data yang dihimpun melalui instrumen OMMHA memerlukan antarmuka web yang dapat membantu tim mengelola catatan survei, melihat status pengajuan, meninjau klasifikasi, menyajikan indikator, serta menampilkan distribusi layanan secara geografis.")
    add_body(doc, "Frontend Dashboard OMMHA menjadi ruang kerja utama untuk kegiatan tersebut. Antarmuka ini menyatukan ringkasan data, manajemen survei, manajemen layanan, analitik, peta, pengguna, sistem, dan dokumentasi bantuan dalam satu navigasi berbasis peran. Laporan ini mendokumentasikan fitur yang tersedia pada versi beta dengan porsi utama pada Manajemen Survei.")
    add_figure(doc, "00-halaman-login.png", "Gambar 1. Halaman masuk Frontend Dashboard OMMHA Versi Beta", width=5.9)

    # 4 — Objective
    new_page_section(doc, "B. Tujuan")
    add_body(doc, "Secara umum, laporan ini bertujuan mendokumentasikan pengembangan dan fungsi antarmuka Frontend Dashboard OMMHA Versi Beta sebagai bagian dari sistem Atlas Layanan Kesehatan Jiwa Kabupaten Kebumen.")
    doc.add_heading("Tujuan khusus", level=2)
    add_numbered(doc, [
        "Mengidentifikasi kelompok fitur yang tersedia melalui navigasi dashboard.",
        "Mendokumentasikan tampilan dan alur kerja frontend untuk pengelolaan survei.",
        "Menjelaskan dukungan frontend bagi pengelolaan layanan, analitik, peta, pengguna, sistem, dan bantuan.",
        "Menyediakan bukti visual dan matriks fitur sebagai bahan komunikasi, demonstrasi, dan serah-terima pengembangan.",
        "Menjelaskan hubungan antarmuka dashboard dengan proses pengumpulan dan pemanfaatan data OMMHA tanpa menguji komponen backend secara terpisah.",
    ])
    doc.add_heading("Pengguna dokumen", level=2)
    add_body(doc, "Dokumen ditujukan bagi Tim Peneliti OMMHA, Pusat Rehabilitasi YAKKUM, pengelola data, administrator, verifier, enumerator, serta pihak lain yang memerlukan gambaran terstruktur mengenai kemampuan frontend dashboard.")

    # 5 — Scope and method
    new_page_section(doc, "C. Ruang Lingkup dan Metode Dokumentasi")
    doc.add_heading("1. Ruang lingkup", level=2)
    add_body(doc, "Cakupan laporan dibatasi pada frontend web yang dapat diakses melalui navigasi dashboard. Backend, API, basis data, dan Mobile Apps OMMHA disebut sebagai konteks integrasi, tetapi tidak menjadi objek pengujian mandiri dalam dokumen ini.")
    add_feature_table(doc, [
        ["Kelompok", "Fokus dokumentasi"],
        ["Dasbor & Analitik", "Ringkasan, indikator, tabel data, laporan, dan ekspor"],
        ["Manajemen Survei", "Daftar, entri, detail, status, unggah massal, model, template, dan audit"],
        ["Manajemen Layanan", "Daftar layanan dan data klasifikasi pendukung"],
        ["Pengguna & Enumerator", "Pengguna, peran, penugasan, kinerja, dan riwayat login"],
        ["Peta & Geospasial", "Lokasi layanan, heatmap, lapisan MTC, wilayah, dan unggah geospasial"],
        ["Sistem dan Bantuan", "Pengaturan, log, panduan, referensi, FAQ, dan dukungan"],
    ], [4.0, 12.0])
    doc.add_heading("2. Metode", level=2)
    add_numbered(doc, [
        "Inventarisasi route dan menu frontend pada repositori aplikasi.",
        "Observasi halaman menggunakan lingkungan pengembangan lokal pada resolusi desktop 1440 × 900 piksel.",
        "Penelusuran navigasi dan elemen interaksi yang terlihat, termasuk filter, pencarian, tabel, status, formulir, visualisasi, dan ekspor.",
        "Pengambilan tangkapan layar aktual pada 20 Juli 2026.",
        "Penyusunan uraian fitur dan matriks dokumentasi berdasarkan kelompok fungsi.",
    ])

    # 6 — Architecture
    new_page_section(doc, "D. Gambaran Umum Frontend Dashboard")
    add_body(doc, "Frontend Dashboard OMMHA menggunakan pola navigasi bertingkat. Rail ikon di sisi paling kiri mengarahkan pengguna ke kelompok modul, sedangkan panel submenu menampilkan fungsi yang relevan. Area kerja di sebelah kanan memuat breadcrumb, judul halaman, kontrol aksi, filter, tabel, kartu indikator, grafik, atau peta sesuai konteks.")
    add_feature_table(doc, [
        ["Lapisan antarmuka", "Peran"],
        ["Autentikasi", "Mengarahkan pengguna masuk ke ruang kerja sesuai sesi yang aktif"],
        ["Navigasi", "Mengelompokkan fitur dalam menu utama dan submenu"],
        ["Presentasi data", "Menampilkan kartu indikator, tabel, grafik, kalender, dan peta"],
        ["Interaksi", "Mendukung pencarian, filter, pengurutan, paginasi, pilihan status, dan aksi halaman"],
        ["Dokumentasi", "Menyediakan panduan pengguna, DESDE-LTC, FAQ, dan dukungan"],
    ], [4.3, 11.7])
    doc.add_heading("Kelompok peran", level=2)
    add_bullets(doc, [
        "Administrator: pengelolaan menyeluruh, termasuk pengguna dan sistem.",
        "Surveyor/Enumerator: entri dan pemantauan survei sesuai penugasan.",
        "Verifier: peninjauan pengajuan dan proses verifikasi.",
        "Viewer: akses baca terhadap informasi, analitik, dan laporan yang diizinkan.",
    ])

    # 7 — Dashboard
    new_page_section(doc, "E. Hasil Dokumentasi Fitur", "1. Ringkasan Dashboard")
    add_body(doc, "Halaman Ringkasan menjadi titik masuk setelah autentikasi. Kartu indikator menyajikan jumlah layanan, survei tertunda, pengguna aktif, dan cakupan kecamatan. Bagian ikhtisar aktivitas menampilkan tren login dan pengajuan survei, sedangkan kalender membantu meninjau catatan berdasarkan tanggal.")
    add_figure(doc, "01-dashboard-ringkasan.png", "Gambar 2. Ringkasan Dashboard OMMHA")
    add_bullets(doc, [
        "Ringkasan indikator utama dalam bentuk kartu statistik.",
        "Visualisasi aktivitas untuk membaca perubahan dari waktu ke waktu.",
        "Kalender dan daftar pengajuan sebagai pintasan pemantauan kegiatan.",
        "Navigasi konsisten menuju modul analitik dan operasional.",
    ])

    # 8 — Indicators
    new_page_section(doc, "2. Indikator Utama")
    add_body(doc, "Halaman Indikator Utama memperluas ringkasan menjadi tampilan analitik. Fitur ini membantu pengguna membaca ketersediaan layanan, distribusi klasifikasi, aktivitas, dan ukuran agregat lain yang disiapkan untuk kebutuhan pemantauan program.")
    add_figure(doc, "02-indikator-utama.png", "Gambar 3. Halaman Indikator Utama")
    add_feature_table(doc, [
        ["Fungsi", "Manfaat"],
        ["Kartu metrik", "Memberikan gambaran cepat atas nilai agregat"],
        ["Grafik distribusi", "Membantu membandingkan kelompok data"],
        ["Filter periode/wilayah", "Mengarahkan pembacaan pada konteks tertentu"],
        ["Penyajian visual", "Memudahkan komunikasi informasi kepada pemangku kepentingan"],
    ], [5.0, 11.0])

    # 9 — Survey management intro
    new_page_section(doc, "F. Manajemen Survei", "Fokus Utama Dokumentasi")
    add_body(doc, "Manajemen Survei merupakan pusat proses pengumpulan dan pemantauan data pada dashboard. Modul ini menghubungkan catatan hasil pengumpulan data dengan status pengajuan, model pertanyaan, template, dan rekam audit. Karena perannya langsung terhadap kualitas serta keterlacakan data OMMHA, bagian ini didokumentasikan lebih rinci dibanding modul lain.")
    add_feature_table(doc, [
        ["Subfitur", "Peran dalam alur survei"],
        ["Semua Catatan Survei", "Menemukan, membandingkan, dan membuka catatan"],
        ["Entri Survei Baru", "Memulai proses pengisian data"],
        ["Detail dan Edit", "Meninjau serta memperbarui isi catatan"],
        ["Pengajuan Tertunda", "Memantau catatan yang menunggu tindak lanjut"],
        ["Unggah Massal", "Memasukkan sejumlah data melalui berkas"],
        ["Model Kuesioner", "Mengelola struktur pertanyaan"],
        ["Template Survei", "Mengatur paket instrumen yang digunakan"],
        ["Log Audit Survei", "Melihat jejak perubahan dan aktivitas"],
    ], [5.0, 11.0])

    # 10 — Survey list
    new_page_section(doc, "1. Semua Catatan Survei")
    add_body(doc, "Daftar survei menyajikan catatan dalam tabel yang dapat dipindai secara cepat. Pengguna dapat melihat status, identitas fasilitas, wilayah, enumerator, serta jawaban kunci yang dipilih sebagai kolom. Kontrol di atas tabel mendukung penyaringan dan pencarian untuk mempercepat penemuan catatan tertentu.")
    add_figure(doc, "03-manajemen-survei.png", "Gambar 4. Daftar Semua Catatan Survei")
    add_bullets(doc, [
        "Filter berdasarkan status, kategori, tanggal, kecamatan, desa, dan enumerator.",
        "Pencarian berdasarkan nama atau informasi yang tersedia.",
        "Pengurutan data dan pemilihan baris.",
        "Penanda status yang membedakan draf, diajukan, terverifikasi, atau ditolak.",
        "Aksi ekspor serta pintasan untuk menambah catatan baru.",
    ])

    # 11 — New survey
    new_page_section(doc, "2. Entri Survei Baru")
    add_body(doc, "Entri Survei Baru digunakan untuk memulai pengumpulan data melalui dashboard web. Halaman mengarahkan pengguna memilih atau mengisi informasi yang diperlukan sesuai instrumen aktif. Struktur bertahap membantu mempertahankan urutan pertanyaan dan konteks jawaban.")
    add_figure(doc, "04-entri-survei-baru.png", "Gambar 5. Halaman awal Entri Survei Baru")
    add_numbered(doc, [
        "Pengguna membuka perintah Entri Survei Baru.",
        "Sistem menampilkan konteks instrumen dan bagian isian yang relevan.",
        "Pengguna melengkapi data secara berurutan serta berpindah antarbagian.",
        "Catatan dapat disimpan sebagai bagian dari proses kerja survei.",
        "Catatan yang tersedia kemudian dapat ditinjau melalui daftar survei.",
    ])

    # 12 — Detail/edit/status
    new_page_section(doc, "3. Detail, Penyuntingan, dan Status Survei")
    add_body(doc, "Setiap catatan survei dapat dibuka untuk melihat informasi lebih lengkap. Tampilan detail mengelompokkan metadata, data fasilitas, jawaban, klasifikasi, dan informasi pendukung agar pengguna memperoleh konteks yang utuh. Ketika izin pengguna memungkinkan, halaman penyuntingan menyediakan mekanisme pembaruan catatan.")
    add_feature_table(doc, [
        ["Tahap", "Dukungan antarmuka"],
        ["Draf", "Catatan masih berada dalam proses pengisian"],
        ["Diajukan", "Catatan dikirim untuk proses berikutnya"],
        ["Terverifikasi", "Status menunjukkan catatan telah melalui verifikasi"],
        ["Ditolak", "Status membedakan catatan yang memerlukan tindak lanjut"],
    ], [4.0, 12.0])
    add_body(doc, "Penanda status digunakan secara konsisten pada daftar dan halaman terkait sehingga pengguna dapat mengenali posisi catatan dalam alur kerja tanpa membuka seluruh isi terlebih dahulu.")

    # 13 — Pending
    new_page_section(doc, "4. Pengajuan Tertunda")
    add_body(doc, "Halaman Pengajuan Tertunda memusatkan catatan yang sedang menunggu verifikasi atau tindak lanjut. Pemisahan ini mendukung pembagian kerja antara pengumpul data dan pengguna yang bertanggung jawab meninjau pengajuan.")
    add_figure(doc, "05-pengajuan-tertunda.png", "Gambar 6. Halaman Pengajuan Tertunda")
    add_bullets(doc, [
        "Daftar terfokus pada catatan berstatus menunggu.",
        "Filter dan pencarian membantu menentukan prioritas peninjauan.",
        "Informasi fasilitas dan wilayah mendukung identifikasi cepat.",
        "Akses menuju detail menjaga kesinambungan dari ringkasan ke pemeriksaan catatan.",
    ])

    # 14 — Bulk upload
    new_page_section(doc, "5. Unggah Massal")
    add_body(doc, "Fitur Unggah Massal menyediakan jalur input melalui berkas ketika data tersedia dalam jumlah banyak. Antarmuka memandu pengguna memilih berkas, memahami format yang diperlukan, dan melanjutkan proses unggah melalui satu halaman khusus.")
    add_figure(doc, "06-unggah-massal.png", "Gambar 7. Fitur Unggah Massal data survei")
    add_numbered(doc, [
        "Menyiapkan data sesuai format yang ditentukan.",
        "Memilih atau menyeret berkas ke area unggah.",
        "Melanjutkan proses melalui kontrol yang tersedia pada halaman.",
        "Memantau hasil melalui catatan survei dan log terkait.",
    ])

    # 15 — Questionnaire model
    new_page_section(doc, "6. Model Kuesioner")
    add_body(doc, "Model Kuesioner menampilkan struktur pertanyaan yang membentuk instrumen survei. Fitur ini penting bagi konsistensi antara kebutuhan penelitian, alur pertanyaan, pilihan jawaban, dan data yang ditampilkan kembali pada tabel hasil.")
    add_figure(doc, "07-model-kuesioner.png", "Gambar 8. Daftar Model Kuesioner OMMHA")
    add_bullets(doc, [
        "Daftar pertanyaan atau model yang tersedia.",
        "Pencarian, filter, dan pengurutan untuk menelusuri struktur instrumen.",
        "Informasi identitas, tipe, urutan, dan status elemen kuesioner.",
        "Akses menuju detail konfigurasi sesuai kewenangan pengguna.",
    ])

    # 16 — Template
    new_page_section(doc, "7. Template Survei")
    add_body(doc, "Template Survei mengelompokkan model pertanyaan menjadi instrumen yang dapat digunakan pada kegiatan pengumpulan data. Halaman template memberi gambaran nama instrumen, versi, status, dan komponen yang menyusunnya.")
    add_figure(doc, "08-template-survei.png", "Gambar 9. Detail Template Survei OMMHA")
    add_feature_table(doc, [
        ["Informasi", "Kegunaan"],
        ["Nama dan versi", "Mengidentifikasi instrumen yang sedang digunakan"],
        ["Status", "Membedakan template aktif dan kondisi lain"],
        ["Daftar bagian/pertanyaan", "Menunjukkan susunan instrumen"],
        ["Metadata", "Mendukung penelusuran konteks pengembangan template"],
    ], [5.0, 11.0])

    # 17 — Audit
    new_page_section(doc, "8. Log Audit Survei")
    add_body(doc, "Log Audit Survei menyediakan catatan aktivitas yang berkaitan dengan perubahan atau pengelolaan survei. Rekam jejak ini mendukung keterlacakan proses dan memudahkan pengguna memahami kapan suatu aktivitas dilakukan serta siapa yang berkaitan dengan aktivitas tersebut.")
    add_figure(doc, "09-log-audit-survei.png", "Gambar 10. Log Audit Survei")
    add_bullets(doc, [
        "Pencatatan waktu aktivitas.",
        "Identifikasi pengguna atau pelaku aktivitas.",
        "Jenis tindakan dan objek survei yang berkaitan.",
        "Filter dan pencarian untuk menelusuri rekam jejak.",
    ])

    # 18 — Survey workflow synthesis
    new_page_section(doc, "9. Alur Terpadu Manajemen Survei")
    add_body(doc, "Berbagai fitur Manajemen Survei membentuk satu alur kerja terpadu. Template dan model kuesioner menjadi dasar instrumen; entri baru menghasilkan catatan; daftar survei menjadi pusat pemantauan; status menggambarkan tahap proses; pengajuan tertunda membantu peninjauan; dan log audit menjaga keterlacakan aktivitas.")
    add_feature_table(doc, [
        ["Urutan", "Tahap", "Fitur frontend"],
        ["1", "Persiapan instrumen", "Model Kuesioner dan Template Survei"],
        ["2", "Pengumpulan data", "Entri Survei Baru atau Unggah Massal"],
        ["3", "Penyimpanan dan pemantauan", "Semua Catatan Survei"],
        ["4", "Pengajuan", "Status Diajukan/Pengajuan Tertunda"],
        ["5", "Peninjauan", "Detail survei dan informasi status"],
        ["6", "Keterlacakan", "Log Audit Survei"],
        ["7", "Pemanfaatan", "Dashboard, analitik, peta, laporan, dan ekspor"],
    ], [1.4, 5.0, 9.6])
    add_body(doc, "Alur tersebut memperlihatkan bahwa Manajemen Survei bukan modul yang berdiri sendiri. Data yang dikelola melalui modul ini menjadi masukan bagi penyajian indikator, tabel layanan, peta, dan laporan pada bagian frontend lainnya.")

    # 19 — Service management
    new_page_section(doc, "G. Modul Pendukung", "1. Manajemen Layanan")
    add_body(doc, "Manajemen Layanan menampilkan data penyedia atau unit layanan yang telah dicatat dalam sistem. Pengguna dapat menelusuri seluruh layanan serta mengakses data referensi yang mendukung klasifikasi DESDE-LTC.")
    add_figure(doc, "10-manajemen-layanan.png", "Gambar 11. Halaman Manajemen Layanan")
    add_bullets(doc, [
        "Semua Layanan dan penambahan layanan baru.",
        "Kategori Layanan atau Basic Stable Inputs of Care (BSIC).",
        "Klasifikasi Main Type of Care (MTC).",
        "Jenis layanan, wilayah geografis, dan populasi sasaran.",
        "Pencarian, filter, pengurutan, dan akses detail data.",
    ])

    # 20 — Data/reporting
    new_page_section(doc, "2. Data, Analitik, dan Laporan")
    add_body(doc, "Kelompok Data dan Analitik membantu pengguna mengubah catatan operasional menjadi informasi yang dapat dibaca dan dibandingkan. Tabel layanan, matriks MTC, analisis cakupan populasi, kesenjangan layanan, laporan ketersediaan, tenaga kerja, serta profil fasilitas tersedia sebagai sudut pandang yang berbeda terhadap data OMMHA.")
    add_feature_table(doc, [
        ["Halaman", "Informasi utama"],
        ["Tabel Layanan", "Daftar terstruktur layanan dan atributnya"],
        ["Matriks MTC", "Perbandingan klasifikasi Main Type of Care"],
        ["Cakupan Populasi", "Gambaran sasaran yang dilayani"],
        ["Kesenjangan Layanan", "Perbandingan ketersediaan antar konteks"],
        ["Laporan Ketersediaan", "Ringkasan ketersediaan layanan"],
        ["Laporan Tenaga Kerja", "Informasi sumber daya manusia"],
        ["Profil Fasilitas", "Gambaran rinci pada tingkat fasilitas"],
    ], [5.0, 11.0])

    # 21 — Map
    new_page_section(doc, "3. Peta & Geospasial")
    add_body(doc, "Peta dan fitur geospasial menyajikan dimensi lokasi dari layanan kesehatan jiwa. Pengguna dapat memindai sebaran titik layanan, mengubah fokus wilayah, dan memanfaatkan lapisan tematik untuk membaca pola spasial.")
    add_figure(doc, "11-peta-layanan.png", "Gambar 12. Peta Lokasi Layanan OMMHA")
    add_bullets(doc, [
        "Peta lokasi layanan dengan titik atau simbol geografis.",
        "Peta panas untuk melihat konsentrasi data.",
        "Lapisan MTC untuk membaca klasifikasi pada peta.",
        "Perbandingan wilayah dan unggah data geospasial.",
        "Kontrol zoom, pencarian, legenda, dan panel informasi sesuai halaman.",
    ])

    # 22 — Users
    new_page_section(doc, "4. Pengguna & Enumerator")
    add_body(doc, "Modul Pengguna dan Enumerator mendukung pengelolaan aktor yang berinteraksi dengan dashboard. Administrator dapat melihat daftar pengguna, peran, dan riwayat login, sedangkan fungsi enumerator membantu pemantauan penugasan serta kinerja pengumpulan data.")
    add_figure(doc, "12-pengguna.png", "Gambar 13. Daftar Pengguna Dashboard OMMHA")
    add_feature_table(doc, [
        ["Fitur", "Dukungan"],
        ["Semua Pengguna", "Daftar akun dan informasi peran"],
        ["Tambah Pengguna", "Formulir pembuatan akun"],
        ["Peran & Izin", "Pengelompokan kewenangan akses"],
        ["Enumerator", "Daftar tenaga pengumpul data"],
        ["Penugasan", "Pengelolaan hubungan enumerator dan kegiatan"],
        ["Kinerja", "Pemantauan aktivitas atau capaian enumerator"],
        ["Riwayat Login", "Jejak akses pengguna"],
    ], [5.0, 11.0])

    # 23 — Export
    new_page_section(doc, "5. Unduh & Ekspor")
    add_body(doc, "Halaman Unduh & Ekspor memberi jalur untuk memperoleh keluaran data dari dashboard. Pengguna memilih cakupan atau jenis informasi melalui kontrol yang tersedia, kemudian menyiapkan keluaran sesuai kebutuhan analisis atau pelaporan lanjutan.")
    add_figure(doc, "13-laporan-ekspor.png", "Gambar 14. Halaman Unduh & Ekspor")
    add_bullets(doc, [
        "Pemilihan jenis data atau laporan.",
        "Penggunaan filter untuk menentukan cakupan keluaran.",
        "Pilihan format yang disediakan oleh antarmuka.",
        "Pemisahan fungsi ekspor dari tampilan operasional harian.",
    ])

    # 24 — System/help
    new_page_section(doc, "6. Sistem, Log, dan Bantuan")
    add_body(doc, "Kelompok Sistem menyediakan pengaturan umum dan log operasional bagi administrator. Log aktivitas, verifikasi, perubahan data, error sistem, serta impor/ekspor membantu pemantauan kegiatan yang berlangsung melalui aplikasi.")
    add_body(doc, "Kelompok Bantuan & Dokumentasi menyediakan panduan pengguna, referensi klasifikasi DESDE-LTC, buku panduan enumerator, FAQ, dan kanal dukungan. Kehadiran materi bantuan di dalam navigasi mengurangi jarak antara pelaksanaan tugas dan kebutuhan mencari penjelasan.")
    add_figure(doc, "14-bantuan.png", "Gambar 15. Panduan Pengguna pada Dashboard OMMHA")
    add_bullets(doc, [
        "Panduan penggunaan dashboard berdasarkan topik.",
        "Referensi konsep dan klasifikasi DESDE-LTC.",
        "Materi praktis bagi enumerator.",
        "Pertanyaan yang sering diajukan dan kontak dukungan.",
    ])

    # 25 — Cross-cutting
    new_page_section(doc, "H. Fitur Lintas Modul")
    add_body(doc, "Selain fungsi khusus pada setiap halaman, frontend dashboard menerapkan sejumlah pola interaksi yang digunakan lintas modul. Pola yang konsisten membantu pengguna memindahkan pengetahuan dari satu halaman ke halaman lain.")
    add_feature_table(doc, [
        ["Fitur lintas modul", "Penerapan"],
        ["Breadcrumb", "Menunjukkan posisi pengguna dalam hierarki halaman"],
        ["Pencarian", "Mempercepat penemuan nama, catatan, atau item tertentu"],
        ["Filter", "Menyaring data menurut status, kategori, waktu, wilayah, atau pengguna"],
        ["Pengurutan", "Mengubah urutan penyajian data"],
        ["Paginasi", "Membagi daftar panjang menjadi halaman yang dapat dikelola"],
        ["Status badge", "Memberikan penanda visual terhadap keadaan data"],
        ["Ekspor", "Menyediakan keluaran data untuk kebutuhan lanjutan"],
        ["Desain responsif", "Menyesuaikan susunan antarmuka dengan ruang layar"],
        ["Navigasi berbasis peran", "Menampilkan fungsi sesuai tanggung jawab pengguna"],
    ], [5.0, 11.0])

    # 26 — Matrix 1
    new_page_section(doc, "Lampiran 1. Matriks Fitur Frontend", "Bagian 1 dari 2")
    add_feature_table(doc, [
        ["No.", "Kelompok", "Halaman/Fitur", "Kegunaan"],
        ["1", "Dasbor", "Ringkasan", "Statistik, aktivitas, kalender, dan pengajuan"],
        ["2", "Dasbor", "Indikator Utama", "Penyajian analitik utama"],
        ["3", "Dasbor", "Pengajuan Terbaru", "Pemantauan catatan terkini"],
        ["4", "Data", "Tabel Layanan", "Daftar layanan terstruktur"],
        ["5", "Data", "Matriks MTC", "Perbandingan klasifikasi MTC"],
        ["6", "Data", "Analisis Cakupan", "Gambaran cakupan populasi"],
        ["7", "Data", "Kesenjangan", "Gambaran kesenjangan layanan"],
        ["8", "Laporan", "Ketersediaan", "Ringkasan ketersediaan layanan"],
        ["9", "Laporan", "Tenaga Kerja", "Informasi sumber daya manusia"],
        ["10", "Laporan", "Profil Fasilitas", "Informasi tingkat fasilitas"],
        ["11", "Laporan", "Unduh & Ekspor", "Keluaran data dan laporan"],
        ["12", "Layanan", "Semua Layanan", "Daftar dan detail layanan"],
        ["13", "Layanan", "Tambah Layanan", "Entri layanan baru"],
        ["14", "Layanan", "Kategori BSIC", "Referensi kategori layanan"],
        ["15", "Layanan", "Klasifikasi MTC", "Referensi Main Type of Care"],
        ["16", "Layanan", "Jenis Layanan", "Referensi jenis layanan"],
        ["17", "Layanan", "Wilayah Geografis", "Referensi wilayah"],
        ["18", "Layanan", "Populasi Sasaran", "Referensi kelompok sasaran"],
    ], [1.0, 2.6, 5.0, 7.4])

    # 27 — Matrix 2
    new_page_section(doc, "Lampiran 1. Matriks Fitur Frontend", "Bagian 2 dari 2")
    add_feature_table(doc, [
        ["No.", "Kelompok", "Halaman/Fitur", "Kegunaan"],
        ["19", "Survei", "Semua Catatan", "Daftar, filter, pencarian, dan ekspor"],
        ["20", "Survei", "Entri Baru", "Pengisian survei melalui web"],
        ["21", "Survei", "Unggah Massal", "Input data melalui berkas"],
        ["22", "Survei", "Pengajuan Tertunda", "Pemantauan catatan menunggu"],
        ["23", "Survei", "Model Kuesioner", "Struktur pertanyaan"],
        ["24", "Survei", "Template", "Paket instrumen dan versi"],
        ["25", "Survei", "Log Audit", "Keterlacakan aktivitas"],
        ["26", "Pengguna", "Semua Pengguna", "Daftar akun dan peran"],
        ["27", "Pengguna", "Enumerator", "Daftar pengumpul data"],
        ["28", "Pengguna", "Penugasan", "Pengelolaan tugas enumerator"],
        ["29", "Pengguna", "Kinerja", "Pemantauan aktivitas enumerator"],
        ["30", "Peta", "Lokasi Layanan", "Sebaran fasilitas pada peta"],
        ["31", "Peta", "Peta Panas", "Konsentrasi spasial"],
        ["32", "Peta", "Lapisan MTC", "Klasifikasi pada peta"],
        ["33", "Peta", "Perbandingan Wilayah", "Pembacaan antarwilayah"],
        ["34", "Sistem", "Pengaturan & Log", "Administrasi dan pemantauan"],
        ["35", "Bantuan", "Panduan & Referensi", "Dukungan penggunaan"],
    ], [1.0, 2.6, 5.0, 7.4])

    # 28 — Conclusion
    new_page_section(doc, "I. Kesimpulan")
    add_body(doc, "Frontend Dashboard OMMHA Versi Beta telah menyediakan ruang kerja terpadu untuk menampilkan, menelusuri, dan mengelola informasi Atlas Layanan Kesehatan Jiwa Kabupaten Kebumen. Struktur navigasi menghubungkan ringkasan dashboard, analitik, layanan, survei, pengguna, peta, sistem, serta bantuan dalam satu antarmuka web.")
    add_body(doc, "Manajemen Survei menjadi bagian paling menonjol karena mencakup rangkaian fungsi dari persiapan instrumen hingga pemanfaatan catatan. Daftar survei, entri baru, detail dan penyuntingan, status pengajuan, unggah massal, model kuesioner, template, serta log audit membentuk alur yang mendukung pengumpulan data secara terstruktur dan dapat ditelusuri.")
    add_body(doc, "Modul pendukung memperluas pemanfaatan data melalui indikator, tabel, peta, laporan, ekspor, pengelolaan pengguna, dan dokumentasi. Dengan demikian, frontend dashboard berfungsi sebagai penghubung antara proses pengumpulan data OMMHA dan kebutuhan membaca informasi untuk pemantauan serta perencanaan layanan kesehatan jiwa.")
    doc.add_paragraph()
    p = doc.add_paragraph("— Akhir laporan —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(TEAL)

    BASE.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
